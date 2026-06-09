from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from openpyxl.styles import PatternFill

MONTH_NAMES = {
    'січень':   1,      'січня':     1,     1:  'січень',      '01': 'січень',
    'лютий':    2,      'лютого':    2,     2:  'лютий',       '02': 'лютий',
    'березень': 3,      'березня':   3,     3:  'березень',    '03': 'березень',
    'квітень':  4,      'квітня':    4,     4:  'квітень',     '04': 'квітень',
    'травень':  5,      'травня':    5,     5:  'травень',     '05': 'травень',
    'червень':  6,      'червня':    6,     6:  'червень',     '06': 'червень',
    'липень':   7,      'липня':     7,     7:  'липень',      '07': 'липень',
    'серпень':  8,      'серпня':    8,     8:  'серпень',     '08': 'серпень',
    'вересень': 9,      'вересня':   9,     9:  'вересень',    '09': 'вересень',
    'жовтень':  10,     'жовтня':    10,    10: 'жовтень',     '10': 'жовтень',
    'листопад': 11,     'листопада': 11,    11: 'листопад',    '11': 'листопад',
    'грудень':  12,     'грудня':    12,    12: 'грудень',     '12': 'грудень'
}

MONTH_NAMES_CAPITAL = {
    1:  'Січень',      '01': 'Січень',
    2:  'Лютий',       '02': 'Лютий',
    3:  'Березень',    '03': 'Березень',
    4:  'Квітень',     '04': 'Квітень',
    5:  'Травень',     '05': 'Травень',
    6:  'Червень',     '06': 'Червень',
    7:  'Липень',      '07': 'Липень',
    8:  'Серпень',     '08': 'Серпень',
    9:  'Вересень',    '09': 'Вересень',
    10: 'Жовтень',     '10': 'Жовтень',
    11: 'Листопад',    '11': 'Листопад',
    12: 'Грудень',     '12': 'Грудень'
}

HOURS_PER_SUBJECT = {
    1: {
        'Біологія':               34,
        'Всесвітня історія':      34,
        'Географія':              40,
        'Зарубіжна література':   34,
        'Захист України':         34,
        'Іноземна мова':          34,
        'Інформатика':            34,
        'Історія України':        34,
        'Математика':             68,
        'Українська література':  34,
        'Українська мова':        34,
        'Фізика':                 68,
        'Фізична культура':       51,
        'Хімія':                  34
    },
    2: {
        'Астрономія':             46,
        'Біологія':               72,
        'Всесвітня історія':      46,
        'Географія':              48,
        'Зарубіжна література':   46,
        'Захист України':         69,
        'Іноземна мова':          46,
        'Інформатика':            23,
        'Історія України':        46,
        'Математика':             69,
        'Українська література':  46,
        'Українська мова':        46,
        'Фізика':                 46,
        'Фізична культура':       69,
        'Хімія':                  54
    },
}

MONTH_NAMES_GEN = [
    'січня',
    'лютого',
    'березня',
    'квітня',
    'травня',
    'червня',
    'липня',
    'серпня',
    'вересня',
    'жовтня',
    'листопада',
    'грудня'
]

EDUCATIONAL_PROGRAMS = {
    'D3':   '«Менеджмент»',
    'D5':   '«Маркетинг»',
    'D7':   '«Підприємництво, торгівля та біржова діяльність»',
    'E2':   '«Екологія»',
    'G1':   '«Хімічні технології та інженерія»',
    'G3':   '«Електроенергетика, електротехніка та електромеханіка»',
    'G5':   "«Комп'ютерні технології та електронні комунікації»",
    'G7':   "«Автоматизація та комп'ютерно-інтегровані технології»",
    'G13':  '«Технологія приготування їжі»',
    'G16':  '«Нафтогазова інженерія та технології»',
    'J2':   '«Готельно-ресторанна справа»',
    'J3':   '«Туризм»'
}

SUBJECT_WITH_2_PROGRAMS = ['G1', 'G5']

FILL_GRAY = PatternFill(start_color='BFBFBF', end_color='BFBFBF', fill_type='solid')

def save_file(file, path):
    for i in range(10):
        # Формируем имя файла
        if i == 0:
            file_path = path
        else:
            base, ext = os.path.splitext(path)
            file_path = f'{base} ({i}){ext}'

        try:
            file.save(file_path)
            if file_path == path:
                return True
            return os.path.basename(file_path)
        except Exception:
            continue  # Игнорируем ошибку и пробуем снова

    raise RuntimeError(f'Не удалось сохранить файл {path}.')

def cleaning_and_save_workbook(workbook, path_to_save):
    answer = None

    # Удаляем лишние листы
    sheets_to_delete = [name for name in workbook.sheetnames if name.startswith('Л')]
    for sheet_name in sheets_to_delete:
        sheet = workbook[sheet_name]
        workbook.remove(sheet)

    # Пытаемся сохранить лист
    saved_workbook_name = save_file(workbook, path_to_save)
    if saved_workbook_name != True:
        answer = saved_workbook_name

    return answer

def replace_text(doc, old_text, new_text):
    for paragraph in doc.paragraphs:
        if old_text not in paragraph.text:
            continue

        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
        # после замены внутри runs можно выйти
    return doc

def explanation_insert(doc, old_text, new_text):
    lines = new_text.split('\n')

    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            # Очищаем весь текст абзаца и заменяем содержимое
            paragraph.text = ''  

            for i, line in enumerate(lines):
                if i == 0:
                    # Первый абзац — тот же, что содержал old_text
                    run = paragraph.add_run(line)
                    p = paragraph
                else:
                    # Следующие строки — новые абзацы
                    new_p = OxmlElement('w:p')
                    paragraph._element.addnext(new_p)
                    p = Paragraph(new_p, paragraph._parent)
                    run = p.add_run(line)

                # Настройка стиля
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.first_line_indent = Cm(1.25)

                # Обновляем "текущий" абзац для корректного добавления следующих
                paragraph = p
            break
    return doc

def short_name(full_name):
    parts = full_name.split()

    # Проверяем, что ровно 3 части
    if len(parts) == 3:
        surname, name, patronymic = parts
        short = f'{surname} {name[0]}.{patronymic[0]}.'
        return short
    else:
        # Если частей не 3 — возвращаем как есть
        return full_name

def insert_row(doc, table_index, text, insert=True, color=False):
    # Получаем список индексов
    table_indices = list(range(len(doc.tables)))

    table = doc.tables[table_index]

    # Добавляем или выбираем последнюю строку
    if insert:
        row = table.add_row()
    else:
        row = table.rows[-1]

    # Номер строки
    row_number = len(table.rows) - 1

    for i, cell in enumerate(row.cells):
        paragraph = cell.paragraphs[0]

        if color:
            # Добавляем заливку ячейки
            cell._tc.get_or_add_tcPr().append(
                parse_xml(f'<w:shd {nsdecls('w')} w:fill="#92D050"/>')
            )

        # Устанавливаем границы ячейки (все стороны)
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_pr.append(parse_xml(
            f"""
            <w:tcBorders {nsdecls('w')}>
                <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>
            </w:tcBorders>
            """
        ))

        # Убираем интервал после абзаца и задаем междустрочный интервал 1.0
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0

        # Устанавливаем высоту строки (0,25 см, минимум)
        row.height = Cm(0.25)

        # Добавляем текст
        if i == 0:
            run = paragraph.add_run(str(row_number))
        else:
            run = paragraph.add_run(str(text[i - 1]) if i - 1 < len(text) else '')

        # Настройка шрифта
        font = run.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        # Центрирование текста
        if i in [0, 2, 3]:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    return doc

def color_table_row(doc, table_index, rowIndex):
    color='#92D050'
    table = doc.tables[table_index]

    fill_color = color.lstrip('#').upper()
    row = table.rows[rowIndex]

    for cell in row.cells:
        # Добавляем/обновляем заливку ячейки
        cell._tc.get_or_add_tcPr().append(
            parse_xml(f'<w:shd {nsdecls('w')} w:fill="{fill_color}"/>')
        )

    return doc

def delete_page(doc, page_to_delete):
    current_page = 1
    elements_to_delete = []
    
    body = doc.element.body
    
    for element in body.iterchildren():
        has_page_break = False
        
        # Проверка на наличие разрыва страницы
        if element.tag.endswith('p'):
            if element.xpath('.//w:br[@w:type="page"]'):
                has_page_break = True
        
        # --- СЦЕНАРИЙ 1: Удаление первой страницы ---
        if page_to_delete == 1:
            if current_page == 1:
                elements_to_delete.append(element)
                if has_page_break:
                    break 

        # --- СЦЕНАРИЙ 2: Удаление страницы N > 1 ---
        else:
            if current_page < page_to_delete - 1:
                if has_page_break:
                    current_page += 1

            elif current_page == page_to_delete - 1:
                if has_page_break:
                    elements_to_delete.append(element)
                    current_page += 1

            elif current_page == page_to_delete:
                if has_page_break:
                    break 
                else:
                    elements_to_delete.append(element)

    if page_to_delete != 1:
        del elements_to_delete[-1]

    for el in elements_to_delete:
        if el.getparent() is not None:
            el.getparent().remove(el)

    return doc

def ukrainian_sort_key(s):
    # Эталонный алфавит (нижний регистр)
    alphabet = 'абвгґдеєжзиіїйклмнопрстуфхцчшщьюя'
    
    result = []
    for char in s.lower():
        index = alphabet.find(char)
        if index != -1:
            result.append(index)
        else:
            # Если символа нет в алфавите (пробел, цифра, спецсимвол), 
            # отправляем его в конец таблицы ASCII
            result.append(ord(char) + 100) 
    return result
