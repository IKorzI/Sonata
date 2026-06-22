from openpyxl.styles import Border, Side, PatternFill, Alignment, Font
from openpyxl.utils import range_boundaries
from openpyxl import load_workbook
from openpyxl.worksheet.worksheet import Worksheet
from openpyxl.utils import get_column_letter
from openpyxl.cell.text import InlineFont
from openpyxl.cell.rich_text import TextBlock, CellRichText
from docx import Document
import os
from pathlib import Path

from utils import (
    MONTH_NAMES_GEN,
    EDUCATIONAL_PROGRAMS,
    SUBJECT_WITH_2_PROGRAMS,
    FILL_GRAY,
)
from utils import (
    save_file,
    delete_page,
    replace_text,
    insert_row,
    short_name,
    explanation_insert,
)


def filling_out_the_summary_sheet(
    sheet: Worksheet,
    students: list,
    cells: list,
    social_len: int,
    subjects: list,
    kurator: str,
    percentage: int,
):
    """
    Filling out the summary sheet.

    Args:
        sheet (Worksheet): The Excel sheet to be filled out.
        students (list): The list of students.
        cells (list): The values of cells "C5", "C6", "C7".
        social_len (int): The number of students in privileged categories.
        subjects (list): The list of subjects.
        kurator (str): The name of the curator.
        percentage (int): The percentage for calculations.
    """

    student_len = len(students)
    subject_len = len(subjects)

    delete_row_start = 9 + student_len + 1
    delete_row_count = 109 - delete_row_start + 1
    end_row = delete_row_start - 1

    delete_col_start = 5 + subject_len + 1
    delete_col_count = 26 - delete_col_start
    end_col = delete_col_start - 1
    end_col_letter = get_column_letter(end_col)
    end_col_plus_1_letter = get_column_letter(end_col + 1)

    dims = sheet.column_dimensions

    # =============================================================================================================
    # PREPARING THE SHEET BEFORE FILLING

    sheet.delete_rows(delete_row_start, amount=delete_row_count)
    for row in range(delete_row_start + 11, delete_row_start + 11 + delete_row_count):
        sheet.row_dimensions[row].hidden = True

    for row in [3, 4, 5, 6, 7]:
        sheet.unmerge_cells(start_row=row, start_column=3, end_row=row, end_column=26)
    sheet.unmerge_cells(start_row=8, start_column=6, end_row=8, end_column=25)
    sheet.unmerge_cells(start_row=8, start_column=26, end_row=9, end_column=26)
    sheet.delete_cols(idx=9, amount=delete_col_count)
    for col in range(delete_col_start + 3, delete_col_start + 3 + delete_col_count):
        dims[get_column_letter(col)].hidden = True
    for row in [3, 4, 5, 6, 7]:
        sheet.merge_cells(
            start_row=row, start_column=3, end_row=row, end_column=delete_col_start
        )
    sheet.merge_cells(
        start_row=8, start_column=6, end_row=8, end_column=delete_col_start - 1
    )
    sheet.merge_cells(
        start_row=8,
        start_column=delete_col_start,
        end_row=9,
        end_column=delete_col_start,
    )
    dims[get_column_letter(delete_col_start)].width = 11.3
    dims[get_column_letter(delete_col_start + 1)].width = 2.7
    dims[get_column_letter(delete_col_start + 2)].width = 2.7

    col_letter = get_column_letter(end_col + 1)
    sheet.print_area = f"C3:{col_letter}{end_row + 9}"

    for col in range(6, end_col + 1):
        col_letter = get_column_letter(col)
        row = end_row + 1
        sheet.cell(row=row, column=col).value = (
            f"=IF(SUM({col_letter}$10:{col_letter}{end_row})=0,"
            f'IF(COUNTA({col_letter}$10:{col_letter}{end_row})=0,"Очікую",0),'
            f"SUM({col_letter}$10:{col_letter}{end_row})/$D{end_row})"
        )
        row = end_row + 2
        sheet.cell(row=row, column=col).value = (
            f"=IF(SUM({col_letter}$10:{col_letter}{end_row})=0,"
            f'IF(COUNTA({col_letter}$10:{col_letter}{end_row})=0,"Очікую",0),'
            f'IF(COUNTIF({col_letter}$10:{col_letter}{end_row},">=4")=0,0,'
            f'COUNTIF({col_letter}$10:{col_letter}{end_row},">=4")/$D{end_row}*100))'
        )
        row = end_row + 3
        sheet.cell(row=row, column=col).value = (
            f"=IF(SUM({col_letter}$10:{col_letter}{end_row})=0,"
            f'IF(COUNTA({col_letter}$10:{col_letter}{end_row})=0,"Очікую",0),'
            f'IF(COUNTIF({col_letter}$10:{col_letter}{end_row},">=7")=0,0,'
            f'COUNTIF({col_letter}$10:{col_letter}{end_row},">=7")/$D{end_row}*100))'
        )
    for row in range(10, end_row + 1):
        sheet.cell(row=row, column=end_col + 1).value = (
            f'=IF(COUNTIF($F{row}:{end_col_letter}{row},"=*")>0,IF($C{row}="К","-"," - "),'
            f'IF(COUNTIF($F{row}:{end_col_letter}{row},"<4")>0,IF(C{row}="К","-"," - "),'
            f'(IFERROR(AVERAGE($F{row}:{end_col_letter}{row}),"Очікую"))))'
        )

    sheet.cell(row=end_row + 1, column=end_col + 1).value = (
        f'=IFERROR(AVERAGE($F{end_row + 1}:{end_col_letter}{end_row + 1}),"Очікую")'
    )

    sheet.cell(row=end_row + 4, column=6).value = (
        f'=IFERROR(AVERAGE(F{end_row + 2}:{end_col_letter}{end_row + 2}),"Очікую")'
    )
    sheet.cell(row=end_row + 5, column=6).value = (
        f'=IFERROR(AVERAGE($F{end_row + 3}:{end_col_letter}{end_row + 3}),"Очікую")'
    )
    sheet.cell(row=end_row + 6, column=6).value = (
        f'=IF(COUNTIF($C$10:$C{end_row},"=Б")=0,IF(COUNTIF($C$10:$C{end_row},"=К")=0,"Очікую",0),COUNTIF($C$10:$C{end_row},"=Б"))'
    )
    sheet.cell(row=end_row + 7, column=6).value = (
        f'=IF(COUNTIF($C$10:$C{end_row},"=К")=0,IF(COUNTIF($C$10:$C{end_row},"=Б")=0,"Очікую",0),COUNTIF($C$10:$C{end_row},"=К"))'
    )
    sheet.cell(row=end_row + 9, column=6).value = (
        f'=IF(F{end_row + 6}="Очікую","Очікую",IF(F{end_row + 6}-COUNTIF(${end_col_plus_1_letter}$10:${end_col_plus_1_letter}{end_row}," - ")<ROUNDDOWN(F{end_row + 6}*{percentage / 100},0),F{end_row + 6}-COUNTIF(${end_col_plus_1_letter}$10:${end_col_plus_1_letter}{end_row}," - "),ROUNDDOWN(F{end_row + 6}*{percentage / 100},0)))'
    )

    # =============================================================================================================
    # FILLING THE SHEET

    sheet.cell(row=5, column=3).value = cells["C5"]
    sheet.cell(row=6, column=3).value = cells["C6"]
    sheet.cell(row=7, column=3).value = cells["C7"]
    sheet.cell(row=end_row + 8, column=6).value = social_len
    if subject_len >= 14:
        sheet.cell(row=end_row + 7, column=16).value = kurator
        sheet.cell(row=end_row + 9, column=16).value = f"Маргарита БРІТІКОВА"
    else:
        sheet.cell(row=end_row + 7, column=21 - delete_col_count).value = kurator
        sheet.cell(row=end_row + 9, column=21 - delete_col_count).value = (
            f"Маргарита БРІТІКОВА"
        )

    for i, subject in enumerate(subjects):
        sheet.cell(row=9, column=i + 6).value = (
            f"{subject['subject_name']}\n{subject['teacher_name']}"
        )

    for i, student in enumerate(students):
        sheet.cell(row=i + 10, column=3).value = student["bc"]
        sheet.cell(row=i + 10, column=5).value = student["student_name"]
        for j, grade in enumerate(student["grades"]):
            sheet.cell(row=i + 10, column=j + 6).value = grade


def filling_out_the_rating_sheet(
    sheet: Worksheet,
    s_title: str,
    scholarship_number: int,
    students: list,
    sorted_list: list,
    subjects: list,
):
    """
    Filling out the rating sheet.

    Args:
        sheet (Worksheet): The Excel sheet to be filled out.
        s_title (str): The name of the summary sheet for references.
        scholarship_number (int): The number of scholarships.
        students (list): The list of students.
        sorted_list (list): The sorted list of student indices.
        subjects (list): The list of subjects.
    """

    student_len = sum(1 for s in students if s["rating_grade"] != None)
    subject_len = len(subjects)

    delete_row_start = 9 + student_len + 1
    delete_row_count = 109 - delete_row_start + 1
    end_row = delete_row_start - 1

    delete_col_start = 5 + subject_len + 1
    delete_col_count = 26 - delete_col_start
    end_col = delete_col_start - 1
    end_col_letter = get_column_letter(end_col)

    dims = sheet.column_dimensions

    # =============================================================================================================
    # PREPARING THE SHEET BEFORE FILLING

    sheet.delete_rows(delete_row_start, amount=delete_row_count)
    for row in range(delete_row_start + 6, delete_row_start + 6 + delete_row_count):
        sheet.row_dimensions[row].hidden = True

    for row in [3, 4, 5, 6, 7]:
        sheet.unmerge_cells(start_row=row, start_column=3, end_row=row, end_column=27)
    sheet.unmerge_cells(start_row=8, start_column=6, end_row=8, end_column=25)
    sheet.unmerge_cells(start_row=8, start_column=26, end_row=9, end_column=26)
    sheet.unmerge_cells(start_row=8, start_column=27, end_row=9, end_column=27)
    sheet.delete_cols(idx=9, amount=delete_col_count)
    for col in range(delete_col_start + 4, delete_col_start + 4 + delete_col_count):
        dims[get_column_letter(col)].hidden = True
    for row in [3, 4, 5, 6, 7]:
        sheet.merge_cells(
            start_row=row, start_column=3, end_row=row, end_column=delete_col_start + 1
        )
    sheet.merge_cells(
        start_row=8, start_column=6, end_row=8, end_column=delete_col_start - 1
    )
    sheet.merge_cells(
        start_row=8,
        start_column=delete_col_start,
        end_row=9,
        end_column=delete_col_start,
    )
    sheet.merge_cells(
        start_row=8,
        start_column=delete_col_start + 1,
        end_row=9,
        end_column=delete_col_start + 1,
    )
    dims[get_column_letter(delete_col_start)].width = 11.3
    dims[get_column_letter(delete_col_start + 1)].width = 11.3
    dims[get_column_letter(delete_col_start + 2)].width = 2.7
    dims[get_column_letter(delete_col_start + 3)].width = 2.7

    col_letter = get_column_letter(end_col + 2)
    sheet.print_area = f"C3:{col_letter}{end_row + 4}"

    for row in range(10, end_row + 1):
        sheet.cell(row=row, column=end_col + 1).value = (
            f'=IF(COUNTIF($F{row}:{end_col_letter}{row},"=*")>0,IF($C{row}="К","-"," - "),'
            f'IF(COUNTIF($F{row}:{end_col_letter}{row},"<4")>0,IF(C{row}="К","-"," - "),'
            f'(IFERROR(AVERAGE($F{row}:{end_col_letter}{row}),"Очікую"))))'
        )
        col_letter = get_column_letter(end_col + 1)
        sheet.cell(row=row, column=end_col + 2).value = (
            f'=IFERROR({col_letter}{row}/12*90,"Очікую")'
        )
    summary_end_row = 9 + len(sorted_list)
    sheet.cell(row=end_row + 1, column=6).value = f"='{s_title}'!F{summary_end_row + 6}"
    sheet.cell(row=end_row + 2, column=6).value = f"='{s_title}'!F{summary_end_row + 7}"
    sheet.cell(row=end_row + 3, column=6).value = f"='{s_title}'!F{summary_end_row + 8}"
    sheet.cell(row=end_row + 4, column=6).value = f"='{s_title}'!F{summary_end_row + 9}"

    # =============================================================================================================
    # FILLING THE SHEET (formulas)

    sheet.cell(row=5, column=3).value = f"='{s_title}'!C5"
    sheet.cell(row=6, column=3).value = f"='{s_title}'!C6"
    sheet.cell(row=7, column=3).value = f"='{s_title}'!C7"
    summary_kurator_row = 9 + len(students) + 7
    if subject_len >= 14:
        kurator_col_letter = get_column_letter(16)
        sheet.cell(row=end_row + 2, column=16).value = (
            f"='{s_title}'!{kurator_col_letter}{summary_kurator_row}"
        )
        sheet.cell(row=end_row + 4, column=16).value = f"Маргарита БРІТІКОВА"
    else:
        kurator_col_letter = get_column_letter(21 - delete_col_count)
        sheet.cell(row=end_row + 2, column=21 - delete_col_count).value = (
            f"='{s_title}'!{kurator_col_letter}{summary_kurator_row}"
        )
        sheet.cell(row=end_row + 4, column=21 - delete_col_count).value = (
            f"Маргарита БРІТІКОВА"
        )

    for i in range(subject_len):
        col = get_column_letter(i + 6)
        sheet.cell(row=9, column=i + 6).value = f"='{s_title}'!{col}9"

    row = 10
    for i in range(student_len):
        index = sorted_list[i]
        sheet.cell(row=row, column=3).value = f"='{s_title}'!C{index + 10}"
        sheet.cell(row=row, column=5).value = f"='{s_title}'!E{index + 10}"
        for j in range(subject_len):
            col = get_column_letter(j + 6)
            sheet.cell(row=row, column=j + 6).value = f"='{s_title}'!{col}{index + 10}"
        row += 1

    for i in range(scholarship_number):
        for j in range(subject_len + 5):
            sheet.cell(row=i + 10, column=3 + j).fill = FILL_GRAY


def session_PackageOfDocuments(info, app_path):
    """
    Forms a package of documents for the session for all necessary specialties.

    Args:
        info (dict): A dictionary with data (specialties, students, curator, semester, etc.).
        app_path (str): The path to the application directory.

    Returns:
        answer (dict): A dictionary with keys 'success' (execution status) and 'files' (paths to generated files).
    """

    path_to_save = Path(info["file_path"]).parent

    answer = {"success": True, "files": []}
    directory_to_save = os.path.dirname(info["file_path"])
    os.makedirs(directory_to_save, exist_ok=True)
    path = f"{app_path}/public/examples/work"

    # =============================================================================================================
    # CREATING A PACKAGE OF DOCUMENTS BY SPECIALTIES

    statements = load_workbook(f"{path}/statements.xlsx")
    step = 0
    for subgroup in info["subgroups"]:
        step += 1

        speciality_code = subgroup["speciality_code"]
        speciality_name = subgroup["speciality_name"]

        socialy_len = sum(
            1 for student in subgroup["students"] if student.get("social_status")
        )

        s_sheet = statements[f"ЛЗ{step}"]
        s_sheet.title = f"Зведена {speciality_code}"
        filling_out_the_summary_sheet(
            s_sheet,
            subgroup["students"],
            subgroup["cells"],
            socialy_len,
            info["subjects"],
            info["kurator_nom"],
            info["percentage"],
        )

        if subgroup["scholarship_number"] == 0:
            continue

        r_sheet = statements[f"ЛР{step}"]
        r_sheet.title = f"Рейтингова {speciality_code}"
        filling_out_the_rating_sheet(
            r_sheet,
            s_sheet.title,
            subgroup["scholarship_number"],
            subgroup["students"],
            subgroup["sorted_list"],
            info["subjects"],
        )

        _semester_start_split = info["semester_start"].split(".")
        _semester_end_split = info["semester_end"].split(".")
        _semester_start_month = MONTH_NAMES_GEN[int(_semester_start_split[1]) - 1]
        _semester_end_month = MONTH_NAMES_GEN[int(_semester_end_split[1]) - 1]
        semester_dates_start = f"«{_semester_start_split[0]}» {_semester_start_month} {_semester_start_split[2]} р."
        semester_dates_end = f"«{_semester_end_split[0]}» {_semester_end_month} {_semester_end_split[2]} р."

        del_increased = len(subgroup["increased_scholarship_list"]) == 0
        del_social = len(subgroup["social_scholarship_list"]) == 0

        doc_petition = Document(f"{path}/petition.docx")
        doc_website = Document(f"{path}/website-rating.docx")
        if del_social:
            doc_petition = delete_page(doc_petition, 3)
        if del_increased:
            doc_petition = delete_page(doc_petition, 1)

        table_index_increased = None if del_increased else 0
        table_index_scholarship = 0 if del_increased else 1
        table_index_social = None if del_social else (1 if del_increased else 2)

        doc_petition = replace_text(doc_petition, "group_code_2", info["group_code"])
        doc_petition = replace_text(
            doc_petition, "speciality_2", f"{speciality_code} {speciality_name}"
        )
        doc_petition = replace_text(doc_petition, "kurator_21", info["kurator_gen"])
        doc_petition = replace_text(doc_petition, "kurator_22", info["kurator_nom"])
        doc_petition = replace_text(
            doc_petition, "semester_number_2", info["semester_number_word"]
        )
        doc_petition = replace_text(doc_petition, "years_2", info["years"])
        doc_petition = replace_text(
            doc_petition, "semester_date_start_2", semester_dates_start
        )
        doc_petition = replace_text(
            doc_petition, "semester_date_end_2", semester_dates_end
        )
        doc_website = replace_text(
            doc_website, "semester_number", info["semester_number"]
        )
        doc_website = replace_text(doc_website, "years", info["years"])
        doc_website = replace_text(
            doc_website, "speciality", f"{speciality_code} {speciality_name}"
        )
        doc_website = replace_text(doc_website, "group_code", info["group_code"])
        doc_website = replace_text(doc_website, "kurator", info["kurator_nom"])
        for i in range(subgroup["scholarship_number"]):
            index = subgroup["sorted_list"][i]
            student = subgroup["students"][index]
            student_name = student["student_name"]
            grade = student["rating_grade"]
            increased = "+" if not del_increased and student["increased"] else ""
            doc_petition = insert_row(
                doc_petition,
                table_index_scholarship,
                [student_name, grade, increased],
                insert=i > 0,
            )
            doc_website = insert_row(
                doc_website,
                0,
                [student_name, grade, increased],
                insert=i > 0,
                color=True,
            )
        for i in range(subgroup["scholarship_number"], len(subgroup["students"])):
            index = subgroup["sorted_list"][i]
            student = subgroup["students"][index]
            if student["rating_grade"] == None:
                continue
            student_name = student["student_name"]
            grade = student["rating_grade"]
            doc_website = insert_row(doc_website, 0, [student_name, grade, ""])

        if not del_social:
            doc_petition = replace_text(
                doc_petition, "group_code_3", info["group_code"]
            )
            doc_petition = replace_text(
                doc_petition, "speciality_3", f"{speciality_code} {speciality_name}"
            )
            doc_petition = replace_text(doc_petition, "kurator_31", info["kurator_gen"])
            doc_petition = replace_text(doc_petition, "kurator_32", info["kurator_nom"])
            doc_petition = replace_text(
                doc_petition, "semester_number_3", info["semester_number_word"]
            )
            doc_petition = replace_text(doc_petition, "years_3", info["years"])
            doc_petition = replace_text(
                doc_petition, "semester_date_start_3", semester_dates_start
            )
            doc_petition = replace_text(
                doc_petition, "semester_date_end_3", semester_dates_end
            )

            for el_index, i in enumerate(subgroup["social_scholarship_list"]):
                student = subgroup["students"][i]
                if student["rating_grade"] == None:
                    continue
                student_name = student["student_name"]
                grade = student["rating_grade"]
                social_status = student["social_status"]
                doc_petition = insert_row(
                    doc_petition,
                    table_index_social,
                    [student_name, grade, social_status],
                    insert=el_index > 0,
                )

        if not del_increased:
            doc_submission = Document(f"{path}/submission.docx")
            doc_petition = replace_text(
                doc_petition, "group_code_1", info["group_code"]
            )
            doc_petition = replace_text(
                doc_petition, "speciality_1", f"{speciality_code} {speciality_name}"
            )
            doc_petition = replace_text(doc_petition, "kurator_11", info["kurator_gen"])
            doc_petition = replace_text(doc_petition, "kurator_12", info["kurator_nom"])
            doc_petition = replace_text(
                doc_petition, "semester_number_1", info["semester_number_word"]
            )
            doc_petition = replace_text(doc_petition, "years_1", info["years"])
            doc_petition = replace_text(
                doc_petition, "semester_date_start_1", semester_dates_start
            )
            doc_petition = replace_text(
                doc_petition, "semester_date_end_1", semester_dates_end
            )
            doc_submission = replace_text(
                doc_submission, "group_code_1", info["group_code"]
            )
            doc_submission = replace_text(
                doc_submission, "group_code_2", info["group_code"]
            )
            doc_submission = replace_text(
                doc_submission, "speciality", f"{speciality_code} {speciality_name}"
            )
            doc_submission = replace_text(
                doc_submission, "kurator_1", info["kurator_gen"]
            )
            doc_submission = replace_text(
                doc_submission, "kurator_2", info["kurator_nom"]
            )
            doc_submission = replace_text(
                doc_submission, "semester_number", info["semester_number_word"]
            )
            doc_submission = replace_text(doc_submission, "years", info["years"])
            submission_text = ""
            for el_index, i in enumerate(subgroup["increased_scholarship_list"]):
                student = subgroup["students"][i]
                student_name = student["student_name"]
                grade = student["rating_grade"]
                doc_petition = insert_row(
                    doc_petition,
                    table_index_increased,
                    [student_name, grade, "+"],
                    insert=el_index > 0,
                )
                submission_text = (
                    f"{short_name(student_name)} –"
                    if el_index == 0
                    else f"{submission_text}\n{short_name(student_name)} –"
                )
            doc_submission = explanation_insert(
                doc_submission, "submission", submission_text
            )

        if len(subgroup["same_scores_list"]) > 0:
            doc_explanation = Document(f"{path}/explanation.docx")
            doc_explanation = replace_text(
                doc_explanation, "group_code_1", info["group_code"]
            )
            doc_explanation = replace_text(
                doc_explanation, "group_code_2", info["group_code"]
            )
            doc_explanation = replace_text(
                doc_explanation, "speciality", f"{speciality_code} {speciality_name}"
            )
            doc_explanation = replace_text(
                doc_explanation, "kurator_1", info["kurator_gen"]
            )
            doc_explanation = replace_text(
                doc_explanation, "kurator_2", info["kurator_nom"]
            )
            doc_explanation = replace_text(
                doc_explanation, "semester_number", info["semester_number_word"]
            )
            doc_explanation = replace_text(doc_explanation, "years", info["years"])
            explanation_text = ""
            for same_score_index, same_score in enumerate(subgroup["same_scores_list"]):
                index = subgroup["sorted_list"][same_score[0]]
                grade = subgroup["students"][index]["rating_grade"]
                same_score_text_part_1 = ""
                same_score_text_part_2 = ""
                for el_index, i in enumerate(same_score):
                    index = subgroup["sorted_list"][i]
                    student_name = short_name(
                        subgroup["students"][index]["student_name"]
                    )
                    if el_index == 0:
                        same_score_text_part_1 = f"{student_name}"
                        same_score_text_part_2 = f"{student_name} завжди бере активну участь у житті групи та допомагає класному керівнику"
                    else:
                        same_score_text_part_1 = (
                            f"{same_score_text_part_1}, {student_name}"
                        )
                        if el_index != len(same_score) - 1:
                            same_score_text_part_2 = f"{same_score_text_part_2}, {student_name} завжди бере активну участь у житті групи та допомагає класному керівнику"
                if same_score_index == 0:
                    explanation_text = (
                        f"{same_score_text_part_1} – {grade}: {same_score_text_part_2}"
                    )
                else:
                    explanation_text = f"{explanation_text}\n{same_score_text_part_1} – {grade}: {same_score_text_part_2}"
            doc_explanation = explanation_insert(
                doc_explanation, "explanation", explanation_text
            )

        doc_petition_path = save_file(
            doc_petition, f"{path_to_save}/Клопотання {speciality_code}.docx"
        )
        if doc_petition_path != True:
            answer["files"].append(doc_petition_path)
        doc_website_path = save_file(
            doc_website, f"{path_to_save}/Рейтинг на сайт {speciality_code}.docx"
        )
        if doc_website_path != True:
            answer["files"].append(doc_website_path)
        if not del_increased:
            doc_submission_path = save_file(
                doc_submission, f"{path_to_save}/Подання {speciality_code}.docx"
            )
            if doc_submission_path != True:
                answer["files"].append(doc_submission_path)
        if len(subgroup["same_scores_list"]) > 0:
            doc_explanation_path = save_file(
                doc_explanation, f"{path_to_save}/Пояснення {speciality_code}.docx"
            )
            if doc_explanation_path != True:
                answer["files"].append(doc_explanation_path)

    statements.remove(statements["Л_Загальна"])
    for i in range(1, 11):
        if f"ЛЗ{i}" in statements.sheetnames:
            statements.remove(statements[f"ЛЗ{i}"])
        if f"ЛР{i}" in statements.sheetnames:
            statements.remove(statements[f"ЛР{i}"])
    workbook_path = save_file(statements, f"{path_to_save}/Відомості.xlsx")
    if workbook_path != True:
        answer["files"].append(workbook_path)

    return answer


def filling_out_the_empty_teacher_statement(
    sheet: Worksheet,
    semester: str,
    year: str,
    students: list,
    subgroups: list,
    group_code: str,
):
    """
    Fills out the empty teacher statement.

    Args:
        sheet (Worksheet): The sheet to be filled out.
        semester (str): The semester (e.g., '1' or 'year').
        year (str): The year of the assessment.
        students (list): The list of students.
        subgroups (list): The list of subgroups (specialties).
        group_code (str): The name or code of the group.
    """

    default_font = InlineFont(rFont="Times New Roman", sz=12)
    underline_font = InlineFont(rFont="Times New Roman", sz=12, u="single")
    underline_bold_font = InlineFont(
        rFont="Times New Roman", sz=12, u="single", b="true"
    )

    merged_cells = [
        [122, 3, 122, 12],
        [123, 3, 123, 8],
        [124, 3, 124, 4],
        [125, 3, 125, 4],
        [126, 3, 126, 4],
        [127, 3, 127, 4],
        [128, 3, 129, 4],
        [124, 5, 124, 6],
        [125, 5, 125, 6],
        [126, 5, 126, 6],
        [127, 5, 127, 6],
        [128, 5, 129, 6],
        [124, 7, 124, 8],
        [125, 7, 125, 8],
        [126, 7, 126, 8],
        [127, 7, 127, 8],
        [128, 7, 129, 8],
        [123, 9, 124, 10],
        [125, 9, 125, 10],
        [126, 9, 126, 10],
        [127, 9, 127, 10],
        [128, 9, 128, 10],
        [129, 9, 129, 10],
        [130, 9, 130, 10],
        [131, 9, 131, 10],
        [132, 9, 132, 10],
        [123, 11, 124, 12],
        [125, 11, 126, 12],
        [127, 11, 127, 12],
        [128, 11, 129, 12],
        [130, 11, 130, 12],
        [131, 11, 131, 12],
        [132, 11, 132, 12],
        [130, 3, 130, 8],
        [131, 3, 131, 8],
        [132, 3, 132, 8],
        [133, 3, 133, 7],
        [133, 8, 133, 12],
        [135, 6, 135, 7],
        [136, 6, 136, 7],
        [137, 6, 137, 7],
        [138, 6, 138, 7],
        [135, 9, 135, 11],
        [136, 9, 136, 11],
        [137, 9, 137, 11],
        [138, 9, 138, 11],
        [140, 3, 140, 12],
        [142, 3, 142, 12],
        [143, 3, 143, 12],
        [144, 3, 144, 12],
    ]

    letter_sizes = {
        "А": 170,
        "а": 105,
        "Б": 134,
        "б": 119,
        "В": 157,
        "в": 110,
        "Г": 133,
        "г": 98,
        "Ґ": 104,
        "ґ": 81,
        "Д": 161,
        "д": 116,
        "Е": 143,
        "е": 105,
        "Ё": 143,
        "ё": 102,
        "Є": 156,
        "є": 99,
        "Э": 155,
        "э": 98,
        "Ж": 212,
        "ж": 162,
        "З": 117,
        "з": 90,
        "И": 172,
        "и": 124,
        "І": 77,
        "і": 60,
        "Ї": 77,
        "ї": 63,
        "Й": 170,
        "й": 124,
        "К": 157,
        "к": 114,
        "Л": 163,
        "л": 116,
        "М": 210,
        "м": 149,
        "Н": 170,
        "н": 124,
        "О": 170,
        "о": 117,
        "П": 170,
        "п": 124,
        "Р": 130,
        "р": 117,
        "С": 157,
        "с": 104,
        "Т": 143,
        "т": 100,
        "У": 170,
        "у": 114,
        "Ф": 184,
        "ф": 152,
        "Х": 172,
        "х": 117,
        "Ц": 170,
        "ц": 124,
        "Ч": 153,
        "ч": 117,
        "Ш": 239,
        "ш": 179,
        "Щ": 239,
        "щ": 182,
        "Ь": 136,
        "ь": 105,
        "ъ": 166,
        "ъ": 121,
        "Ю": 246,
        "ю": 175,
        "Я": 156,
        "я": 107,
        "A": 170,
        "a": 105,
        "B": 157,
        "b": 116,
        "C": 157,
        "c": 105,
        "D": 170,
        "d": 117,
        "E": 143,
        "e": 103,
        "F": 129,
        "f": 78,
        "G": 172,
        "g": 114,
        "H": 170,
        "h": 117,
        "I": 76,
        "i": 62,
        "J": 92,
        "j": 63,
        "K": 170,
        "k": 117,
        "L": 138,
        "l": 64,
        "M": 210,
        "m": 183,
        "N": 172,
        "n": 117,
        "O": 170,
        "o": 117,
        "P": 129,
        "p": 116,
        "Q": 170,
        "q": 117,
        "R": 157,
        "r": 79,
        "S": 127,
        "s": 90,
        "T": 144,
        "t": 63,
        "U": 170,
        "u": 117,
        "V": 170,
        "v": 114,
        "W": 225,
        "w": 170,
        "X": 170,
        "x": 114,
        "Y": 170,
        "y": 114,
        "Z": 141,
        "z": 100,
        "0": 114,
        "1": 117,
        "2": 117,
        "3": 116,
        "4": 117,
        "5": 117,
        "6": 116,
        "7": 117,
        "8": 117,
        "9": 116,
        " ": 59,
        "_": 124,
        ",": 57,
        ".": 59,
        ";": 61,
        ":": 63,
        "'": 40,
        "`": 73,
        '"': 100,
        "-": 76,
        "−": 132,
        "–": 117,
        "—": 237,
        "=": 132,
        "+": 132,
        "÷": 129,
        "/": 61,
        "\\": 61,
        "(": 76,
        ")": 77,
        "{": 112,
        "}": 109,
        "[": 81,
        "]": 78,
        "*": 116,
        "№": 226,
        "@": 218,
        "#": 117,
        "$": 117,
        "%": 196,
        "^": 110,
        "&": 183,
        "?": 110,
        "!": 76,
        "~": 126,
        "<": 132,
        ">": 132,
        "«": 110,
        "»": 109,
    }

    # =============================================================================================================
    # SPECIALTIES AND EDUCATIONAL PROGRAMS

    line_size = 11299
    # "Specialty "         1448
    # "Educational program "  1825

    max_specs_width = line_size - 1448 - 200
    max_educational_programs_width = line_size - 1825 - 200

    def get_width(text):
        return sum(letter_sizes.get(char, 100) for char in text)

    comma_width = get_width(", ")
    space_width = get_width(" ")

    # Helper function to split text into lines of a given width
    def append_to_lines(items, max_width):
        lines = []
        current_line = ""
        current_line_width = 0

        for i, item in enumerate(items):
            item_width = get_width(item)
            needs_comma = i > 0
            prefix = ", " if needs_comma else ""
            prefix_width = comma_width if needs_comma else 0

            if current_line_width + prefix_width + item_width <= max_width:
                current_line += prefix + item
                current_line_width += prefix_width + item_width
            else:
                if needs_comma and current_line:
                    current_line += ", "
                    current_line_width += comma_width

                words = item.split()
                for j, word in enumerate(words):
                    word_width = get_width(word)
                    word_prefix = " " if j > 0 else ""
                    word_prefix_width = space_width if j > 0 else 0

                    if not current_line:
                        current_line = word
                        current_line_width = word_width
                    elif (
                        current_line_width + word_prefix_width + word_width <= max_width
                    ):
                        current_line += word_prefix + word
                        current_line_width += word_prefix_width + word_width
                    else:
                        lines.append(current_line)
                        current_line = word
                        current_line_width = word_width

        if current_line:
            lines.append(current_line)

        padding = " " * 200
        return [line + padding for line in lines]

    spec_strings = [
        f"{subgroup['speciality_code']} {subgroup['speciality_name']}"
        for subgroup in subgroups
    ]
    edu_strings = [
        EDUCATIONAL_PROGRAMS[subgroup["speciality_code"]]
        for subgroup in subgroups
        if EDUCATIONAL_PROGRAMS.get(subgroup["speciality_code"], "")
    ]

    specs_text = append_to_lines(spec_strings, max_specs_width)
    educational_programs_text = append_to_lines(
        edu_strings, max_educational_programs_width
    )

    # =============================================================================================================
    # PREPARING THE SHEET BEFORE FILLING

    start = len(students) + 20 + 1
    delete_count = 120 - start + 1

    target_range = "C124:L146"
    min_col, min_row, max_col, max_row = range_boundaries(target_range)
    to_unmerge = []
    for merged in sheet.merged_cells.ranges:
        if (
            merged.min_row >= min_row
            and merged.max_row <= max_row
            and merged.min_col >= min_col
            and merged.max_col <= max_col
        ):
            to_unmerge.append(merged)
    for cell_range in to_unmerge:
        sheet.unmerge_cells(str(cell_range))

    sheet.delete_rows(start, amount=delete_count)

    for row in range(start + 25 + 1, start + 25 + 1 + delete_count):
        sheet.row_dimensions[row].hidden = True

    target_range = f"C{start + 1}:L{start + 1 + 22}"
    min_col, min_row, max_col, max_row = range_boundaries(target_range)
    to_unmerge = []
    for merged in sheet.merged_cells.ranges:
        if (
            merged.min_row >= min_row
            and merged.max_row <= max_row
            and merged.min_col >= min_col
            and merged.max_col <= max_col
        ):
            to_unmerge.append(merged)
    for cell_range in to_unmerge:
        sheet.unmerge_cells(str(cell_range))

    for mc in merged_cells:
        sheet.merge_cells(
            start_row=mc[0] - delete_count,
            start_column=mc[1],
            end_row=mc[2] - delete_count,
            end_column=mc[3],
        )

    cell_range = f"G{start + 2}:L{start + 2 + 9}"
    thin_side = Side(border_style="thin", color="000000")
    white_fill = PatternFill(
        start_color="FFFFFF", end_color="FFFFFF", fill_type="solid"
    )
    center_align = Alignment(wrap_text=True, horizontal="center", vertical="center")
    top_align = Alignment(wrap_text=True, horizontal="center", vertical="top")
    border_style = Border(
        left=thin_side, right=thin_side, top=thin_side, bottom=thin_side
    )
    rows = sheet[cell_range]
    for row in rows:
        for cell in row:
            cell.border = border_style
            cell.fill = white_fill
            cell.alignment = center_align
            cell.font = Font(name="Times New Roman", size=12)

    border_style = Border(bottom=thin_side)

    sheet.cell(row=start + 14, column=6).border = border_style
    sheet.cell(row=start + 16, column=6).border = border_style

    for col in [6, 7, 9, 10, 11]:
        sheet.cell(row=start + 14, column=col).border = border_style
        sheet.cell(row=start + 16, column=col).border = border_style

    for col in [9, 10, 11]:
        sheet.cell(row=start + 14, column=col).font = Font(
            name="Times New Roman", size=12
        )
        sheet.cell(row=start + 14, column=col).alignment = center_align
        sheet.cell(row=start + 14, column=col).fill = white_fill
        sheet.cell(row=start + 15, column=col).font = Font(
            name="Times New Roman", size=10
        )
        sheet.cell(row=start + 15, column=col).alignment = top_align
        sheet.cell(row=start + 15, column=col).fill = white_fill
        sheet.cell(row=start + 16, column=col).font = Font(
            name="Times New Roman", size=12
        )
        sheet.cell(row=start + 16, column=col).alignment = center_align
        sheet.cell(row=start + 16, column=col).fill = white_fill
        sheet.cell(row=start + 17, column=col).font = Font(
            name="Times New Roman", size=10
        )
        sheet.cell(row=start + 17, column=col).alignment = top_align
        sheet.cell(row=start + 17, column=col).fill = white_fill

    for col in [6, 7]:
        sheet.cell(row=start + 15, column=col).font = Font(
            name="Times New Roman", size=10
        )
        sheet.cell(row=start + 15, column=col).alignment = top_align
        sheet.cell(row=start + 15, column=col).fill = white_fill
        sheet.cell(row=start + 17, column=col).font = Font(
            name="Times New Roman", size=10
        )
        sheet.cell(row=start + 17, column=col).alignment = top_align
        sheet.cell(row=start + 17, column=col).fill = white_fill

    border_style = Border(bottom=thin_side)
    for col in [8, 9, 10, 11]:
        sheet.cell(row=start + 12, column=col).font = Font(
            name="Times New Roman", size=12
        )
        sheet.cell(row=start + 12, column=col).alignment = center_align
        sheet.cell(row=start + 12, column=col).fill = white_fill
        sheet.cell(row=start + 12, column=col).border = border_style
    border_style = Border(bottom=thin_side, right=thin_side)
    sheet.cell(row=start + 12, column=12).font = Font(name="Times New Roman", size=12)
    sheet.cell(row=start + 12, column=12).alignment = center_align
    sheet.cell(row=start + 12, column=12).fill = white_fill
    sheet.cell(row=start + 12, column=12).border = border_style
    sheet.cell(row=start + 12, column=12).fill = white_fill

    for row in range(start + 13, start + 20 + 1):
        sheet.cell(row=row, column=8).fill = white_fill

    sheet.row_dimensions[start + 1].height = 25
    sheet.row_dimensions[start + 3].height = 30
    sheet.row_dimensions[start + 13].height = 20
    sheet.row_dimensions[start + 15].height = 20
    sheet.row_dimensions[start + 18].height = 20
    sheet.row_dimensions[start + 20].height = 20
    sheet.row_dimensions[start + 23].height = 21

    # =============================================================================================================
    # FILLING THE SHEET

    sheet.cell(row=10, column=3).value = CellRichText(
        TextBlock(default_font, "Курс "),
        TextBlock(underline_font, " I "),
        TextBlock(default_font, "     Група "),
        TextBlock(underline_font, f" {group_code} "),
    )

    sheet.cell(row=10, column=12).value = CellRichText(
        TextBlock(default_font, "Навчальний рік "),
        TextBlock(underline_font, " 2025-2026 "),
    )

    if semester == "рік":
        sheet.cell(row=14, column=3).value = CellRichText(
            TextBlock(default_font, "за "), TextBlock(underline_font, " рік ")
        )
    else:
        sheet.cell(row=14, column=3).value = CellRichText(
            TextBlock(default_font, "за "),
            TextBlock(underline_font, f" {semester} "),
            TextBlock(default_font, " навчальний семестр"),
        )

    sheet.cell(row=15, column=3).value = CellRichText(
        TextBlock(default_font, "Форма семестрового контролю "),
        TextBlock(underline_bold_font, " залік "),
    )

    sheet.cell(row=15, column=12).value = CellRichText(
        TextBlock(default_font, f"Дата «____» __________ {year}р.")
    )

    for i, el in enumerate(students):
        sheet.cell(row=i + 21, column=4).value = el["student_name"]

    sheet.cell(row=start + 2, column=9).value = "ВСЬОГО\nОЦІНОК"
    sheet.cell(row=start + 2, column=11).value = "Державний\nнорматив ***"
    sheet.cell(row=start + 4, column=11).value = "Не менше 50"
    sheet.cell(row=start + 7, column=11).value = "Не більше 10%"

    sheet.cell(row=start + 15, column=9).value = "(ім'я та прізвище)"
    sheet.cell(row=start + 15, column=6).value = "(підпис)"
    sheet.cell(row=start + 16, column=9).value = "Маргарита БРІТІКОВА"
    sheet.cell(row=start + 17, column=9).value = "(ім'я та прізвище)"
    sheet.cell(row=start + 17, column=6).value = "(підпис)"
    sheet.cell(row=start + 19, column=3).value = f"«____» _______ _____ {year} р."

    sheet.cell(row=start + 12, column=8).value = "Якісна успішність –        %"

    if len(specs_text) == 1:
        sheet.cell(row=6, column=3).value = CellRichText(
            TextBlock(default_font, "Спеціальність "),
            TextBlock(underline_font, specs_text[0]),
        )
        sheet.unmerge_cells(start_row=6, start_column=3, end_row=6, end_column=12)
        sheet.unmerge_cells(start_row=7, start_column=3, end_row=7, end_column=12)
        sheet.merge_cells(start_row=6, start_column=3, end_row=7, end_column=12)
        sheet.row_dimensions[6].height = 7.5
        sheet.row_dimensions[7].height = 7.5
    else:
        sheet.cell(row=6, column=3).value = CellRichText(
            TextBlock(default_font, "Спеціальність "),
            TextBlock(underline_font, specs_text[0]),
        )
        sheet.cell(row=7, column=3).value = CellRichText(
            TextBlock(underline_font, specs_text[1])
        )

    if len(educational_programs_text) == 1:
        sheet.cell(row=8, column=3).value = CellRichText(
            TextBlock(default_font, "Освітня програма "),
            TextBlock(underline_font, educational_programs_text[0]),
        )
        sheet.unmerge_cells(start_row=8, start_column=3, end_row=8, end_column=12)
        sheet.unmerge_cells(start_row=9, start_column=3, end_row=9, end_column=12)
        sheet.merge_cells(start_row=8, start_column=3, end_row=9, end_column=12)
        sheet.row_dimensions[8].height = 7.5
        sheet.row_dimensions[9].height = 7.5
    else:
        sheet.cell(row=8, column=3).value = CellRichText(
            TextBlock(default_font, "Освітня програма "),
            TextBlock(underline_font, educational_programs_text[0]),
        )
        sheet.cell(row=9, column=3).value = CellRichText(
            TextBlock(underline_font, educational_programs_text[1])
        )

    i = 0
    if len(specs_text) != 1:
        i += 1
    if len(educational_programs_text) != 1:
        i += 1

    studLen = len(students)
    if studLen > 35 - i:
        sheet.print_area = f"C3:L{20 + studLen + 2 + 22}"
    elif studLen > 21 - i:
        sheet.print_area = (
            f"C3:L{20 + studLen},C{20 + studLen + 2}:L{20 + studLen + 2 + 22}"
        )
    else:
        sheet.print_area = f"C3:L{20 + studLen + 2 + 22}"


def filling_out_the_empty_sheet(
    sheet: Worksheet,
    cells: list,
    students: list,
    subject_len: int,
    last_student_row: int,
    percentage: int,
):
    """
    Fills out an empty grade statement sheet for a subgroup.

    Args:
        sheet (Worksheet): The sheet that needs to be filled.
        cells (list): A list of values for the header cells (e.g., C5, C6).
        students (list): A list of students in the subgroup.
        subject_len (int): The number of subjects.
        last_student_row (int): The index of the last row containing a student on the general sheet.
        percentage (int): The percentage for performance calculations.
    """

    student_len = len(students)

    delete_row_start = 9 + student_len + 1
    delete_row_count = 109 - delete_row_start + 1
    end_row = delete_row_start - 1

    delete_col_start = 5 + subject_len + 1
    delete_col_count = 26 - delete_col_start
    end_col = delete_col_start - 1
    end_col_letter = get_column_letter(end_col)
    end_col_plus_1_letter = get_column_letter(end_col + 1)

    dims = sheet.column_dimensions

    # =============================================================================================================
    # PREPARING THE SHEET BEFORE FILLING

    sheet.delete_rows(delete_row_start, amount=delete_row_count)
    for row in range(delete_row_start + 11, delete_row_start + 11 + delete_row_count):
        sheet.row_dimensions[row].hidden = True

    for row in [3, 4, 5, 6, 7]:
        sheet.unmerge_cells(start_row=row, start_column=3, end_row=row, end_column=26)
    sheet.unmerge_cells(start_row=8, start_column=6, end_row=8, end_column=25)
    sheet.unmerge_cells(start_row=8, start_column=26, end_row=9, end_column=26)
    sheet.delete_cols(idx=9, amount=delete_col_count)
    for col in range(delete_col_start + 3, delete_col_start + 3 + delete_col_count):
        dims[get_column_letter(col)].hidden = True
    for row in [3, 4, 5, 6, 7]:
        sheet.merge_cells(
            start_row=row, start_column=3, end_row=row, end_column=delete_col_start
        )
    sheet.merge_cells(
        start_row=8, start_column=6, end_row=8, end_column=delete_col_start - 1
    )
    sheet.merge_cells(
        start_row=8,
        start_column=delete_col_start,
        end_row=9,
        end_column=delete_col_start,
    )
    dims[get_column_letter(delete_col_start)].width = 11.3
    dims[get_column_letter(delete_col_start + 1)].width = 2.7
    dims[get_column_letter(delete_col_start + 2)].width = 2.7

    col_letter = get_column_letter(end_col)
    sheet.print_area = f"C3:{col_letter}{end_row + 9}"

    for col in range(6, end_col + 1):
        col_letter = get_column_letter(col)
        row = end_row + 1
        sheet.cell(row=row, column=col).value = (
            f"=IF(SUM({col_letter}$10:{col_letter}{end_row})=0,"
            f'IF(COUNTA({col_letter}$10:{col_letter}{end_row})=0,"Очікую",0),'
            f"SUM({col_letter}$10:{col_letter}{end_row})/$D{end_row})"
        )
        row = end_row + 2
        sheet.cell(row=row, column=col).value = (
            f"=IF(SUM({col_letter}$10:{col_letter}{end_row})=0,"
            f'IF(COUNTA({col_letter}$10:{col_letter}{end_row})=0,"Очікую",0),'
            f'IF(COUNTIF({col_letter}$10:{col_letter}{end_row},">=4")=0,0,'
            f'COUNTIF({col_letter}$10:{col_letter}{end_row},">=4")/$D{end_row}*100))'
        )
        row = end_row + 3
        sheet.cell(row=row, column=col).value = (
            f"=IF(SUM({col_letter}$10:{col_letter}{end_row})=0,"
            f'IF(COUNTA({col_letter}$10:{col_letter}{end_row})=0,"Очікую",0),'
            f'IF(COUNTIF({col_letter}$10:{col_letter}{end_row},">=7")=0,0,'
            f'COUNTIF({col_letter}$10:{col_letter}{end_row},">=7")/$D{end_row}*100))'
        )
    for row in range(10, end_row + 1):
        sheet.cell(row=row, column=end_col + 1).value = (
            f'=IF(COUNTIF($F{row}:{end_col_letter}{row},"=*")>0,IF($C{row}="K","-"," - "),'
            f'IF(COUNTIF($F{row}:{end_col_letter}{row},"<4")>0,IF(C{row}="K","-"," - "),'
            f'(IFERROR(AVERAGE($F{row}:{end_col_letter}{row}),"Очікую"))))'
        )
    sheet.cell(row=end_row + 1, column=end_col + 1).value = (
        f'=IFERROR(AVERAGE($F{end_row + 1}:{end_col_letter}{end_row + 1}),"Очікую")'
    )

    sheet.cell(row=end_row + 4, column=6).value = (
        f'=IFERROR(AVERAGE(F{end_row + 2}:{end_col_letter}{end_row + 2}),"Очікую")'
    )
    sheet.cell(row=end_row + 5, column=6).value = (
        f'=IFERROR(AVERAGE($F{end_row + 3}:{end_col_letter}{end_row + 3}),"Очікую")'
    )
    sheet.cell(row=end_row + 6, column=6).value = (
        f'=IF(COUNTIF($C$10:$C{end_row},"=B")=0,IF(COUNTIF($C$10:$C{end_row},"=K")=0,"Очікую",0),COUNTIF($C$10:$C{end_row},"=B"))'
    )
    sheet.cell(row=end_row + 7, column=6).value = (
        f'=IF(COUNTIF($C$10:$C{end_row},"=K")=0,IF(COUNTIF($C$10:$C{end_row},"=B")=0,"Очікую",0),COUNTIF($C$10:$C{end_row},"=K"))'
    )
    sheet.cell(row=end_row + 9, column=6).value = (
        f'=IF(F{end_row + 6}="Очікую","Очікую",IF(F{end_row + 6}-COUNTIF(${end_col_plus_1_letter}$10:${end_col_plus_1_letter}{end_row}," - ")<ROUNDDOWN(F{end_row + 6}*{percentage / 100},0),F{end_row + 6}-COUNTIF(${end_col_plus_1_letter}$10:${end_col_plus_1_letter}{end_row}," - "),ROUNDDOWN(F{end_row + 6}*{percentage / 100},0)))'
    )

    # =============================================================================================================
    # FILLING THE SHEET

    sheet.cell(row=5, column=3).value = cells["C5"]
    sheet.cell(row=6, column=3).value = cells["C6"]
    sheet.cell(row=7, column=3).value = cells["C7"]
    if subject_len >= 14:
        kurator_col_letter = get_column_letter(16)
        sheet.cell(row=end_row + 7, column=16).value = (
            f"='Загальна'!{kurator_col_letter}{last_student_row + 7}"
        )
        sheet.cell(row=end_row + 9, column=16).value = f"Маргарита БРІТІКОВА"
    else:
        kurator_col_letter = get_column_letter(21 - delete_col_count)
        sheet.cell(row=end_row + 7, column=21 - delete_col_count).value = (
            f"='Загальна'!{kurator_col_letter}{last_student_row + 7}"
        )
        sheet.cell(row=end_row + 9, column=21 - delete_col_count).value = (
            f"Маргарита БРІТІКОВА"
        )

    start_col = 6
    end_col = 6 + subject_len
    for col in range(start_col, end_col):
        fcol_letter = get_column_letter(col)
        sheet.cell(row=9, column=col).value = f"='Загальна'!{fcol_letter}9"

    for el_index, student_index in enumerate(students):
        row = el_index + 10
        frow = student_index + 10
        sheet.cell(row=row, column=3).value = f"='Загальна'!C{frow}"
        sheet.cell(row=row, column=5).value = f"='Загальна'!E{frow}"
        for col in range(start_col, end_col):
            fcol_letter = get_column_letter(col)
            sheet.cell(row=row, column=col).value = f"='Загальна'!{fcol_letter}{frow}"


def filling_out_the_general_empty_sheet(
    sheet: Worksheet, students: list, subjects: list, kurator: str, percentage: int
):
    """
    Fills out the general empty grade statement sheet for the entire group.

    Args:
        sheet (Worksheet): The sheet that needs to be filled.
        students (list): A list of all students in the group.
        subjects (list): The general list of subjects.
        kurator (str): The full name of the group curator.
        percentage (int): The percentage for calculations (e.g., for scholarships).
    """

    student_len = len(students)
    subject_len = len(subjects)

    delete_row_start = 9 + student_len + 1
    delete_row_count = 109 - delete_row_start + 1
    end_row = delete_row_start - 1

    delete_col_start = 5 + subject_len + 1
    delete_col_count = 26 - delete_col_start
    end_col = delete_col_start - 1
    end_col_letter = get_column_letter(end_col)
    end_col_plus_1_letter = get_column_letter(end_col + 1)

    dims = sheet.column_dimensions

    # =============================================================================================================
    # PREPARING THE SHEET BEFORE FILLING

    sheet.delete_rows(delete_row_start, amount=delete_row_count)
    for row in range(delete_row_start + 11, delete_row_start + 11 + delete_row_count):
        sheet.row_dimensions[row].hidden = True

    sheet.unmerge_cells(start_row=3, start_column=3, end_row=7, end_column=26)
    sheet.unmerge_cells(start_row=8, start_column=6, end_row=8, end_column=25)
    sheet.unmerge_cells(start_row=8, start_column=26, end_row=9, end_column=26)
    sheet.delete_cols(idx=9, amount=delete_col_count)
    for col in range(delete_col_start + 3, delete_col_start + 3 + delete_col_count):
        dims[get_column_letter(col)].hidden = True
    sheet.merge_cells(
        start_row=3, start_column=3, end_row=7, end_column=delete_col_start
    )
    sheet.merge_cells(
        start_row=8, start_column=6, end_row=8, end_column=delete_col_start - 1
    )
    sheet.merge_cells(
        start_row=8,
        start_column=delete_col_start,
        end_row=9,
        end_column=delete_col_start,
    )
    dims[get_column_letter(delete_col_start)].width = 11.3
    dims[get_column_letter(delete_col_start + 1)].width = 2.7
    dims[get_column_letter(delete_col_start + 2)].width = 2.7

    col_letter = get_column_letter(end_col)
    sheet.print_area = f"C3:{col_letter}{end_row + 9}"

    for col in range(6, end_col + 1):
        col_letter = get_column_letter(col)
        row = end_row + 1
        sheet.cell(row=row, column=col).value = (
            f"=IF(SUM({col_letter}$10:{col_letter}{end_row})=0,"
            f'IF(COUNTA({col_letter}$10:{col_letter}{end_row})=0,"Очікую",0),'
            f"SUM({col_letter}$10:{col_letter}{end_row})/$D{end_row})"
        )
        row = end_row + 2
        sheet.cell(row=row, column=col).value = (
            f"=IF(SUM({col_letter}$10:{col_letter}{end_row})=0,"
            f'IF(COUNTA({col_letter}$10:{col_letter}{end_row})=0,"Очікую",0),'
            f'IF(COUNTIF({col_letter}$10:{col_letter}{end_row},">=4")=0,0,'
            f'COUNTIF({col_letter}$10:{col_letter}{end_row},">=4")/$D{end_row}*100))'
        )
        row = end_row + 3
        sheet.cell(row=row, column=col).value = (
            f"=IF(SUM({col_letter}$10:{col_letter}{end_row})=0,"
            f'IF(COUNTA({col_letter}$10:{col_letter}{end_row})=0,"Очікую",0),'
            f'IF(COUNTIF({col_letter}$10:{col_letter}{end_row},">=7")=0,0,'
            f'COUNTIF({col_letter}$10:{col_letter}{end_row},">=7")/$D{end_row}*100))'
        )
    for row in range(10, end_row + 1):
        sheet.cell(row=row, column=end_col + 1).value = (
            f'=IF(COUNTIF($F{row}:{end_col_letter}{row},"=*")>0,IF($C{row}="K","-"," - "),'
            f'IF(COUNTIF($F{row}:{end_col_letter}{row},"<4")>0,IF($C{row}="K","-"," - "),'
            f'(IFERROR(AVERAGE($F{row}:{end_col_letter}{row}),"Очікую"))))'
        )
    sheet.cell(row=end_row + 1, column=end_col + 1).value = (
        f'=IFERROR(AVERAGE($F{end_row + 1}:{end_col_letter}{end_row + 1}),"Очікую")'
    )

    sheet.cell(row=end_row + 4, column=6).value = (
        f'=IFERROR(AVERAGE(F{end_row + 2}:{end_col_letter}{end_row + 2}),"Очікую")'
    )
    sheet.cell(row=end_row + 5, column=6).value = (
        f'=IFERROR(AVERAGE($F{end_row + 3}:{end_col_letter}{end_row + 3}),"Очікую")'
    )
    sheet.cell(row=end_row + 6, column=6).value = (
        f'=IF(COUNTIF($C$10:$C{end_row},"=B")=0,IF(COUNTIF($C$10:$C{end_row},"=K")=0,"Очікую",0),COUNTIF($C$10:$C{end_row},"=B"))'
    )
    sheet.cell(row=end_row + 7, column=6).value = (
        f'=IF(COUNTIF($C$10:$C{end_row},"=K")=0,IF(COUNTIF($C$10:$C{end_row},"=B")=0,"Очікую",0),COUNTIF($C$10:$C{end_row},"=K"))'
    )
    sheet.cell(row=end_row + 9, column=6).value = (
        f'=IF(F{end_row + 6}="Очікую","Очікую",IF(F{end_row + 6}-COUNTIF(${end_col_plus_1_letter}$10:${end_col_plus_1_letter}{end_row}," - ")<ROUNDDOWN(F{end_row + 6}*{percentage / 100},0),F{end_row + 6}-COUNTIF(${end_col_plus_1_letter}$10:${end_col_plus_1_letter}{end_row}," - "),ROUNDDOWN(F{end_row + 6}*{percentage / 100},0)))'
    )

    # =============================================================================================================
    # FILLING THE SHEET

    if subject_len >= 14:
        sheet.cell(row=end_row + 7, column=16).value = kurator
        sheet.cell(row=end_row + 9, column=16).value = f"Маргарита БРІТІКОВА"
    else:
        sheet.cell(row=end_row + 7, column=21 - delete_col_count).value = kurator
        sheet.cell(row=end_row + 9, column=21 - delete_col_count).value = (
            f"Маргарита БРІТІКОВА"
        )

    for i, subj in enumerate(subjects):
        sheet.cell(row=9, column=i + 6).value = (
            f"{subj['subject_name']}\n{subj['teacher_name']}"
        )

    for i, student in enumerate(students):
        sheet.cell(row=i + 10, column=3).value = student["bc"]
        sheet.cell(row=i + 10, column=5).value = student["student_name"]


def filling_out_the_journal_sheet(
    sheet: Worksheet, subjects: list, start_index: int, group_name: str
):
    """
    Fills out a journal sheet with a list of subjects.

    Args:
        sheet (Worksheet): The sheet that needs to be filled.
        subjects (list): A list of subjects.
        start_index (int): The starting index for numbering.
        group_name (str): The name of the group.
    """

    subjects_len = len(subjects)

    delete_row_start = 4 + subjects_len + 1
    delete_row_count = 24 - delete_row_start + 1
    end_row = delete_row_start - 1

    # =============================================================================================================
    # PREPARING THE SHEET BEFORE FILLING

    sheet.delete_rows(delete_row_start, amount=delete_row_count)
    for row in range(delete_row_start + 2, delete_row_start + 3 + delete_row_count):
        sheet.row_dimensions[row].hidden = True

    col_letter = get_column_letter(10)
    sheet.print_area = f"C3:{col_letter}{end_row}"

    # =============================================================================================================
    # FILLING THE SHEET

    sheet.cell(row=3, column=3).value = group_name

    for index, subject in enumerate(subjects):
        row = 5 + index
        sheet.cell(row=row, column=4).value = start_index + index
        sheet.cell(row=row, column=6).value = subject["subject_name"]
        sheet.cell(row=row, column=7).value = subject["teacher_name"]


def session_EmptyCreate(info, app_path, path_to_save, semester, subcject_index=None):
    """
    Creates empty grade statements and journals for the session.

    Args:
        info (dict): Dictionary with information about groups, students, and subjects.
        app_path (str): Path to the application directory.
        path_to_save (str): Path to save the generated files.
        semester (str): Semester ('I', 'II', or 'year').
        subcject_index (int, optional): Starting index for subject numbering.
    Returns:
        list/dict: List of saved files or a dictionary with the subject index and files (for the II semester).
    """
    files = []
    path = f"{app_path}/public/examples/work/"

    # Loading document templates
    journal = load_workbook(f"{path}/journal.xlsx")
    teacher_statements = load_workbook(f"{path}/teacher-statements.xlsx")

    # Configuring the journal title page
    sheet = journal[f"Титульна"]
    if semester == "рік":
        text = f"рік"
    else:
        text = f"{semester} семестр"
    sheet.cell(row=9, column=3).value = text
    sheet.cell(row=11, column=3).value = f"{info['years']} н.р."

    group_index = -1
    if not subcject_index:
        subcject_index = info["first_index"]

    # Processing each group
    for group in info["groups"]:
        group_index += 1
        subject_len = len(group["subjects"])
        group_code = group["group_code"]

        sheet = journal[f"Л{group_index + 1}"]
        sheet.title = group_code
        filling_out_the_journal_sheet(
            sheet, group["subjects"], subcject_index, group_code
        )
        subcject_index += subject_len

        # Loading and filling the general statement
        statements = load_workbook(f"{path}/statements.xlsx")
        sheet = statements[f"Л_Загальна"]
        sheet.title = "Загальна"
        filling_out_the_general_empty_sheet(
            sheet,
            group["students"],
            group["subjects"],
            group["kurator_nom"],
            info["percentage"],
        )

        # Filling statements for each subgroup
        for subgroup_index, subgroup in enumerate(group["subgroups"]):
            cells = {
                "C5": f"Успішності студентів спеціальності {subgroup['speciality_code']} {subgroup['speciality_name']}",
                "C6": (
                    f"За {info['semester_roman']} семестр {info['years']} н.р."
                    if semester != "рік"
                    else f"За рік {info['years']} н.р."
                ),
                "C7": f"Група {group_code} курс I",
            }
            sheet = statements[f"ЛЗ{subgroup_index + 1}"]
            sheet.title = f"Зведена {subgroup['speciality_code']}"
            last_student_row = 9 + len(group["students"])
            filling_out_the_empty_sheet(
                sheet,
                cells,
                subgroup["student_IDs"],
                subject_len,
                last_student_row,
                info["percentage"],
            )

        # Filling statements for teachers
        sheet = teacher_statements[f"Л{group_index + 1}"]
        sheet.title = f"{group_code}"
        subgroups = []
        for subgroup in group["subgroups"]:
            subgroups.append(
                {
                    "speciality_code": subgroup["speciality_code"],
                    "speciality_name": subgroup["speciality_name"],
                }
            )
        filling_out_the_empty_teacher_statement(
            sheet, semester, info["year"], group["students"], subgroups, group_code
        )

        # Creating a directory to save group files
        pathToGroupSave = os.path.join(path_to_save, group_code)
        os.makedirs(pathToGroupSave, exist_ok=True)

        # Deleting unnecessary template sheets before saving
        for sheet_name in statements.sheetnames:
            if sheet_name and sheet_name[0] in ["Л"]:
                del statements[sheet_name]

        workbook_path = save_file(
            statements, f"{pathToGroupSave}/Відомості {group_code}.xlsx"
        )
        if workbook_path != True:
            files.append(workbook_path)

    # Deleting unnecessary templates in the journal
    for sheet_name in journal.sheetnames:
        if sheet_name and sheet_name[0] in ["Л"]:
            del journal[sheet_name]

    workbook_path = save_file(journal, f"{path_to_save}/Журнал видачі відомостей.xlsx")
    if workbook_path != True:
        files.append(workbook_path)

    # Deleting unnecessary templates in teacher statements
    for sheet_name in teacher_statements.sheetnames:
        if sheet_name and sheet_name[0] in ["Л"]:
            del teacher_statements[sheet_name]

    workbook_path = save_file(
        teacher_statements, f"{path_to_save}/Відомості для викладачів.xlsx"
    )
    if workbook_path != True:
        files.append(workbook_path)

    # Returning results depending on the current semester
    if semester == "II":
        return {"subcject_index": subcject_index, "files": files}
    else:
        return files


def session_EmptyStart(info, app_path):
    """
    Entry point for starting the generation of empty session documents.

    Args:
        info (dict): Main dictionary with information about groups, semester, etc.
        app_path (str): Path to the application directory.
    Returns:
        dict: Dictionary with execution status, generated files, and possible warning text.
    """
    answer = {"success": True, "files": [], "customText": ""}
    parent_dir = Path(info["file_path"]).parent
    path_to_save = None

    # Forming a warning text if there are specialties with multiple programs
    for group in info["groups"]:
        for subgroup in group["subgroups"]:
            if subgroup["speciality_code"] in SUBJECT_WITH_2_PROGRAMS:
                if answer["customText"] == "":
                    answer["customText"] = "{{python.session.emptyStart.customText}}"
                answer["customText"] = (
                    f"{answer['customText']}\n  –  [{group['group_code']}] {subgroup['speciality_code']}: {EDUCATIONAL_PROGRAMS[subgroup['speciality_code']]}"
                )

    # Generation logic for the I semester
    if info["semester_number"] == 1:
        path_to_save = os.path.join(parent_dir, "I семестр")
        os.makedirs(path_to_save, exist_ok=True)
        files = session_EmptyCreate(info, app_path, path_to_save, "I")
        if files:
            for file in files:
                answer["files"].append(file)

    # Generation logic for the II semester and annual statements
    else:
        path_to_save = os.path.join(parent_dir, "II семестр")
        os.makedirs(path_to_save, exist_ok=True)
        result = session_EmptyCreate(info, app_path, path_to_save, "II")
        files = result["files"]
        subcject_index = result["subcject_index"]
        if files:
            for file in files:
                answer["files"].append(file)

        path_to_save = os.path.join(parent_dir, "рік")
        os.makedirs(path_to_save, exist_ok=True)
        files = session_EmptyCreate(info, app_path, path_to_save, "рік", subcject_index)
        if files:
            for file in files:
                answer["files"].append(file)

    return answer


def session_ReportStart(info, app_path):
    """
    Forms and saves a general report (PZSO) based on group data.

    Args:
        info (dict): Dictionary with information about groups (performance statistics, budget/contract, etc.).
        app_path (str): Path to the application root directory.
    Returns:
        answer (dict): Dictionary with execution status ('success') and list of saving errors ('files').
    """
    path_to_save = Path(info["file_path"]).parent

    answer = {"success": True, "files": []}
    directory_to_save = os.path.dirname(info["file_path"])
    os.makedirs(directory_to_save, exist_ok=True)
    path = f"{app_path}/public/examples/work"

    # =============================================================================================================
    # REPORT CREATION

    report = load_workbook(f"{path}/report.xlsx")
    sheet = report["Дані"]
    step = 0
    general_amount = 0

    # Iterating through each group to populate statistics
    for group in info["groups"]:
        row = 6 + step
        general_amount += group["amount"]

        sheet.cell(row=row, column=2).value = group["group_code"]

        sheet.cell(row=row, column=4).value = group["budget"]
        sheet.cell(row=row, column=5).value = group["kontrakt"]

        sheet.cell(row=row, column=6).value = group["scholarship"]
        sheet.cell(row=row, column=7).value = group["social_scholarship"]

        sheet.cell(row=row, column=8).value = group["amount"]
        sheet.cell(row=row, column=9).value = 0

        # Filling performance data (high, sufficient, middle, low levels)
        sheet.cell(row=row, column=12).value = group["achievement"]["hight"]
        sheet.cell(row=row, column=13).value = group["achievement"]["sufficient"]
        sheet.cell(row=row, column=14).value = group["achievement"]["middle"]
        sheet.cell(row=row, column=15).value = group["achievement"]["low"]

        sheet.cell(row=row, column=19).value = group["avg_grade"]

        step += 1

    # Recording the total number of students
    sheet.cell(row=4, column=22).value = general_amount

    doc_path = save_file(report, f"{path_to_save}/ПЗСО.xlsx")
    if doc_path != True:
        answer["files"].append(doc_path)

    return answer


def session_DebtorsStart(info, app_path):
    path_to_save = Path(info["file_path"]).parent

    answer = {"success": True, "files": []}
    directory_to_save = os.path.dirname(info["file_path"])
    os.makedirs(directory_to_save, exist_ok=True)
    path = f"{app_path}/public/examples/work"

    # =============================================================================================================
    # CREATING A TABLE OF DEBTORS

    debtors = load_workbook(f"{path}/debtors.xlsx")

    step = 0
    for group in info["groups"]:
        step += 1

        debrost_sheet = debtors[f"Л{step}"]
        debrost_sheet.title = group["group_code"]
        debrost_sheet.cell(row=4, column=3).value = group["group_code"]

        debrost_row = 5
        debrost_step = 1
        for student in group["students"]:
            subject_step = 0
            student_row = 0
            for grade in student["grades"]:
                if student_row == 0:
                    debrost_sheet.cell(row=debrost_row, column=3).value = debrost_step
                    debrost_sheet.cell(row=debrost_row, column=4).value = student["bc"]
                    debrost_sheet.cell(row=debrost_row, column=5).value = student[
                        "student_name"
                    ]
                    student_row = 1
                elif student_row == 1:
                    debrost_sheet.merge_cells(
                        start_row=debrost_row - 1,
                        start_column=3,
                        end_row=debrost_row,
                        end_column=3,
                    )
                    debrost_sheet.merge_cells(
                        start_row=debrost_row - 1,
                        start_column=4,
                        end_row=debrost_row,
                        end_column=4,
                    )
                    debrost_sheet.merge_cells(
                        start_row=debrost_row - 1,
                        start_column=5,
                        end_row=debrost_row,
                        end_column=5,
                    )
                    student_row = 2
                else:
                    debrost_sheet.unmerge_cells(
                        start_row=debrost_row - student_row,
                        start_column=3,
                        end_row=debrost_row - 1,
                        end_column=3,
                    )
                    debrost_sheet.unmerge_cells(
                        start_row=debrost_row - student_row,
                        start_column=4,
                        end_row=debrost_row - 1,
                        end_column=4,
                    )
                    debrost_sheet.unmerge_cells(
                        start_row=debrost_row - student_row,
                        start_column=5,
                        end_row=debrost_row - 1,
                        end_column=5,
                    )
                    debrost_sheet.merge_cells(
                        start_row=debrost_row - student_row,
                        start_column=3,
                        end_row=debrost_row,
                        end_column=3,
                    )
                    debrost_sheet.merge_cells(
                        start_row=debrost_row - student_row,
                        start_column=4,
                        end_row=debrost_row,
                        end_column=4,
                    )
                    debrost_sheet.merge_cells(
                        start_row=debrost_row - student_row,
                        start_column=5,
                        end_row=debrost_row,
                        end_column=5,
                    )
                    student_row += 1
                debrost_sheet.cell(row=debrost_row, column=6).value = grade["grade"]
                debrost_sheet.cell(row=debrost_row, column=7).value = grade[
                    "subject_name"
                ]
                debrost_sheet.cell(row=debrost_row, column=8).value = grade[
                    "teacher_name"
                ]
                debrost_row += 1
                subject_step += 1
            debrost_step += 1

        delete_row_start = debrost_row
        delete_row_count = 104 - debrost_row + 1
        debrost_sheet.delete_rows(delete_row_start, amount=delete_row_count)
        for row in range(delete_row_start + 2, delete_row_start + 2 + delete_row_count):
            debrost_sheet.row_dimensions[row].hidden = True
        debrost_sheet.row_dimensions[delete_row_start].height = 14.1
        debrost_sheet.row_dimensions[delete_row_start + 1].height = 14.1

        debrost_sheet.print_area = f"C3:H{delete_row_start - 1}"

    for i in range(1, 21):
        if f"Л{i}" in debtors.sheetnames:
            debtors.remove(debtors[f"Л{i}"])

    doc_path = save_file(debtors, f"{path_to_save}/Боржники.xlsx")
    if doc_path != True:
        answer["files"].append(doc_path)

    return answer
